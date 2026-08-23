import pandas as pd
import json
import re
import os
import sys
import urllib.request
import io
import shutil
from docx import Document

sys.stdout.reconfigure(encoding='utf-8')

SHEET_ID = "1kXlVrNZ9wEGPTrW5ClvGcoi037Oad8rk"

XLSX_URL = f"https://docs.google.com/spreadsheets/d/{SHEET_ID}/export?format=xlsx"
CSV_URLS = [
    f"https://docs.google.com/spreadsheets/d/{SHEET_ID}/export?format=csv",
    f"https://docs.google.com/spreadsheets/d/{SHEET_ID}/gviz/tq?tqx=out:csv",
    f"https://docs.google.com/spreadsheets/d/{SHEET_ID}/pub?output=csv"
]

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
json_path = os.path.join(BASE_DIR, 'meetings_db.json')
html_path = os.path.join(BASE_DIR, 'index.html')
local_excel_path = os.path.join(BASE_DIR, 'extracted_agendas_live.xlsx')
package_dir = os.path.join(BASE_DIR, 'BTRC_Meeting_Portal_Package')

print("🌐 1/4 Fetching live data from Google Sheet...")

df = None

# 1. Try XLSX first (preserves full structured workbook)
try:
    req = urllib.request.Request(XLSX_URL, headers={'User-Agent': 'Mozilla/5.0'})
    with urllib.request.urlopen(req, timeout=20) as response:
        content = response.read()
        df = pd.read_excel(io.BytesIO(content))
        print(f"  ✓ Successfully fetched full Google Sheet (XLSX format) with {len(df)} rows.")
except Exception as e:
    print(f"  Notice: XLSX fetch fallback ({e})")

# 2. Fallback to CSV if needed
if df is None or len(df) < 5:
    for url in CSV_URLS:
        try:
            req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
            with urllib.request.urlopen(req, timeout=15) as response:
                content = response.read()
                df = pd.read_csv(io.BytesIO(content))
                if len(df) > 5:
                    print(f"  ✓ Successfully fetched Google Sheet (CSV format) from: {url}")
                    break
        except Exception:
            pass

# 3. Fallback to local excel
if df is None:
    print("  ⚠️ Warning: Could not fetch Google Sheet live. Falling back to local Excel database...")
    if os.path.exists(local_excel_path):
        df = pd.read_excel(local_excel_path)
    else:
        raise RuntimeError("No data source available!")

df = df.fillna('')

# Fuzzy column finder
def find_col(df, keywords):
    for col in df.columns:
        col_str = str(col).strip()
        for kw in keywords:
            if kw in col_str:
                return col
    return None

m_no_col = find_col(df, ['মিটিং নম্বর', 'কমিশন সভা', 'Meeting No'])
m_date_col = find_col(df, ['তারিখ', 'Date'])
a_no_col = find_col(df, ['এজেন্ডা নম্বর', 'এজেন্ডা নং', 'Agenda'])
subj_col = find_col(df, ['বিষয়', 'Subject'])
dec_col = find_col(df, ['সিদ্ধান্ত', 'Decision'])
fine_col = find_col(df, ['জরিমানা', 'Fine'])
impl_col = find_col(df, ['বাস্তবায়ন', 'দপ্তর', 'বিভাগ', 'Implementation'])
status_col = find_col(df, ['অবস্থা', 'স্ট্যাটাস', 'Status'])

meetings_dict = {}

for idx, row in df.iterrows():
    m_no = str(row.get(m_no_col, '') if m_no_col else '').strip()
    m_date = str(row.get(m_date_col, '') if m_date_col else '').strip()
    a_no = str(row.get(a_no_col, '') if a_no_col else '').strip()
    subject = str(row.get(subj_col, '') if subj_col else '').strip()
    decision = str(row.get(dec_col, '') if dec_col else '').strip()
    fine = str(row.get(fine_col, '') if fine_col else '').strip()
    impl = str(row.get(impl_col, '') if impl_col else '').strip()
    status = str(row.get(status_col, '') if status_col else '').strip()

    if not m_no or m_no == 'nan':
        continue

    # Ensure format like "২৭৮তম"
    if not ('তম' in m_no or 'th' in m_no.lower() or 'মিটিং' in m_no):
        m_no = f"{m_no}তম"

    if m_no not in meetings_dict:
        meetings_dict[m_no] = {
            "meeting_number": m_no,
            "meeting_date": m_date,
            "agendas": []
        }

    meetings_dict[m_no]["agendas"].append({
        "agenda_no": a_no if a_no and a_no != 'nan' else str(len(meetings_dict[m_no]["agendas"]) + 1),
        "subject": subject if subject != 'nan' else '',
        "decision": decision if decision != 'nan' else '',
        "fine_amount": fine if fine != 'nan' else '',
        "details": {
            "presentation_summary": f"বিষয়: {subject}" if subject != 'nan' else '',
            "tables": [],
            "implementation": impl if impl and impl != 'nan' else "এনফোর্সমেন্ট এন্ড ইন্সপেকশন ডিরেক্টরেট",
            "assigned_inspector": "",
            "case_status": status if status and status != 'nan' else "চলমান"
        }
    })

def meeting_sort_key(m):
    num_str = re.sub(r'\D', '', m["meeting_number"])
    return int(num_str) if num_str.isdigit() else 999

sorted_meetings = sorted(list(meetings_dict.values()), key=meeting_sort_key, reverse=True)

coordination_mom = []
commissioner_instructions = []
if os.path.exists(json_path):
    try:
        with open(json_path, 'r', encoding='utf-8') as f:
            old_db = json.load(f)
            coordination_mom = old_db.get('coordination_mom', [])
            commissioner_instructions = old_db.get('commissioner_instructions', [])
    except Exception:
        pass

master_db_object = {
    "meetings": sorted_meetings,
    "coordination_mom": coordination_mom,
    "commissioner_instructions": commissioner_instructions
}

total_agendas = sum(len(m['agendas']) for m in sorted_meetings)
print(f"💾 2/4 Saving JSON Database with {len(sorted_meetings)} meetings ({total_agendas} agendas)...")
with open(json_path, 'w', encoding='utf-8') as f:
    json.dump(master_db_object, f, ensure_ascii=False, indent=2)

print("🌐 3/4 Updating Web Portal (index.html)...")
with open(html_path, 'r', encoding='utf-8') as f:
    html = f.read()

json_str = json.dumps(master_db_object, ensure_ascii=False, indent=4)

pdf_map_path = os.path.join(BASE_DIR, 'meeting_drive_pdf_map.json')
pdf_map_str = "{}"
if os.path.exists(pdf_map_path):
    with open(pdf_map_path, 'r', encoding='utf-8') as pf:
        pdf_map_str = pf.read().strip()

pdf_inject_code = f"""
        const gdriveMainFolderUrl = "https://drive.google.com/drive/folders/1mgWMNrL24N92irRgrd4WUL3hhYrfSsIj?usp=sharing";
        const meetingPdfDriveMap = {pdf_map_str};

        function getMeetingPdfUrl(meetingNumber) {{
            if (!meetingNumber) return gdriveMainFolderUrl;
            let numStr = meetingNumber.toString().replace(/[০-৯]/g, d => banglaDigits[d]).replace(/\\D/g, '');
            if (numStr && meetingPdfDriveMap[numStr]) {{
                return meetingPdfDriveMap[numStr];
            }}
            return gdriveMainFolderUrl;
        }}
"""

start_idx = html.find('const masterDb = {')
end_idx = html.find('const banglaDigits =', start_idx)
if start_idx != -1 and end_idx != -1:
    substring = html[start_idx:end_idx]
    last_brace_idx = substring.rfind('};')
    new_text = f'const masterDb = {json_str};\n{pdf_inject_code}\n        '
    new_html = html[:start_idx] + new_text + html[start_idx + last_brace_idx + 2:]

    with open(html_path, 'w', encoding='utf-8') as f:
        f.write(new_html)

# Update local excel backup
try:
    df.to_excel(local_excel_path, index=False)
except Exception as e:
    print(f"  Notice: Could not update Excel backup ({e})")

print("📝 4/4 Updating MS Word Documents (cm word copy)...")
word_dir = os.path.join(BASE_DIR, 'cm word copy')
os.makedirs(word_dir, exist_ok=True)
try:
    for f_name in os.listdir(word_dir):
        if f_name.endswith('.docx'):
            try:
                os.remove(os.path.join(word_dir, f_name))
            except Exception:
                pass

    for meeting in sorted_meetings:
        meeting_no = meeting.get('meeting_number', 'Unknown')
        safe_filename = re.sub(r'[\\/*?:"<>|]', '_', meeting_no)
        filename = os.path.join(word_dir, f"{safe_filename}.docx")
        
        doc = Document()
        doc.add_heading(f"কমিশন সভা: {meeting_no}", 0)
        doc.add_paragraph(f"তারিখ: {meeting.get('meeting_date', '')}")
        
        for a in meeting.get('agendas', []):
            agenda_no = a.get('agenda_no', '')
            doc.add_heading(f"এজেন্ডা {agenda_no}", level=1)
            
            doc.add_heading("বিষয় (Subject):", level=2)
            doc.add_paragraph(a.get('subject', ''))
            
            doc.add_heading("সিদ্ধান্ত (Decision):", level=2)
            doc.add_paragraph(a.get('decision', ''))
            
            if a.get('fine_amount'):
                doc.add_paragraph(f"প্রশাসনিক জরিমানা: {a.get('fine_amount')} টাকা")
                
            doc.add_paragraph("-" * 40)
            
        doc.save(filename)
except Exception as err:
    print(f"  Notice: Word document generation skipped ({err})")

# Also sync to package directory if present
if os.path.exists(package_dir):
    try:
        shutil.copy2(json_path, os.path.join(package_dir, 'meetings_db.json'))
        shutil.copy2(html_path, os.path.join(package_dir, 'index.html'))
        if os.path.exists(local_excel_path):
            shutil.copy2(local_excel_path, os.path.join(package_dir, 'extracted_agendas_live.xlsx'))
        pkg_word_dir = os.path.join(package_dir, 'cm word copy')
        if os.path.exists(word_dir):
            shutil.copytree(word_dir, pkg_word_dir, dirs_exist_ok=True)
        shutil.copy2(__file__, os.path.join(package_dir, 'sync_from_google_sheet.py'))
        print("  ✓ Package directory synchronized.")
    except Exception as e:
        print(f"  Notice: Package copy skipped ({e})")

print(f"🎉 Synchronization Complete! ({len(sorted_meetings)} Meetings, {total_agendas} Agendas)")
