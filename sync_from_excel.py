import pandas as pd
import json
import re
import os
import sys
from docx import Document

sys.stdout.reconfigure(encoding='utf-8')

excel_path = r'e:\user\OneDrive - Bangladesh Telecommunication Regulatory Commission\1. E & I\Meeting\extracted_agendas_live.xlsx'
json_path = r'e:\user\OneDrive - Bangladesh Telecommunication Regulatory Commission\1. E & I\Meeting\meetings_db.json'
html_path = r'e:\user\OneDrive - Bangladesh Telecommunication Regulatory Commission\1. E & I\Meeting\index.html'

print("📊 1/4 Reading Master Excel File (extracted_agendas_live.xlsx)...")
df = pd.read_excel(excel_path)

# Fill NaNs with empty string
df = df.fillna('')

meetings_dict = {}

for idx, row in df.iterrows():
    m_no = str(row.get('মিটিং নম্বর', '')).strip()
    m_date = str(row.get('মিটিংয়ের তারিখ', '')).strip()
    a_no = str(row.get('এজেন্ডা নম্বর', '')).strip()
    subject = str(row.get('বিষয়', '')).strip()
    decision = str(row.get('সিদ্ধান্ত', '')).strip()
    fine = str(row.get('প্রশাসনিক জরিমানা (টাকা)', '')).strip()
    impl = str(row.get('বাস্তবায়নকারী বিভাগ', '')).strip()
    status = str(row.get('মামলার অবস্থা', '')).strip()

    if not m_no:
        continue

    if m_no not in meetings_dict:
        meetings_dict[m_no] = {
            "meeting_number": m_no,
            "meeting_date": m_date,
            "agendas": []
        }

    meetings_dict[m_no]["agendas"].append({
        "agenda_no": a_no if a_no else str(len(meetings_dict[m_no]["agendas"]) + 1),
        "subject": subject,
        "decision": decision,
        "fine_amount": fine,
        "details": {
            "presentation_summary": f"বিষয়: {subject}",
            "tables": [],
            "implementation": impl if impl else "এনফোর্সমেন্ট এন্ড ইন্সপেকশন ডিরেক্টরেট",
            "assigned_inspector": "",
            "case_status": status if status else "চলমান"
        }
    })

# Convert dict to sorted meetings list
def meeting_sort_key(m):
    num_str = re.sub(r'\D', '', m["meeting_number"])
    return int(num_str) if num_str.isdigit() else 999

sorted_meetings = sorted(list(meetings_dict.values()), key=meeting_sort_key, reverse=True)

# Preserve coordination_mom & commissioner_instructions from existing db if any
coordination_mom = []
commissioner_instructions = []
if os.path.exists(json_path):
    try:
        with open(json_path, 'r', encoding='utf-8') as f:
            old_db = json.load(f)
            coordination_mom = old_db.get('coordination_mom', [])
            commissioner_instructions = old_db.get('commissioner_instructions', [])
    except Exception:
        pass

master_db_object = {
    "meetings": sorted_meetings,
    "coordination_mom": coordination_mom,
    "commissioner_instructions": commissioner_instructions
}

# Save updated meetings_db.json
print("💾 2/4 Saving JSON Database (meetings_db.json)...")
with open(json_path, 'w', encoding='utf-8') as f:
    json.dump(master_db_object, f, ensure_ascii=False, indent=2)

# Update index.html
print("🌐 3/4 Updating Web Portal (index.html)...")
with open(html_path, 'r', encoding='utf-8') as f:
    html = f.read()

json_str = json.dumps(master_db_object, ensure_ascii=False, indent=4)

pdf_map_path = os.path.join(BASE_DIR, 'meeting_drive_pdf_map.json')
pdf_map_str = "{}"
if os.path.exists(pdf_map_path):
    with open(pdf_map_path, 'r', encoding='utf-8') as pf:
        pdf_map_str = pf.read().strip()

pdf_inject_code = f"""
        let meetings = masterDb.meetings;
        let coordinationMom = masterDb.coordination_mom || [];
        let commissionerInstructions = masterDb.commissioner_instructions || [];
        let activeMeetingIndex = null;
        let activeAgendaIndex = null;
        let viewMode = 'list';

        const gdriveMainFolderUrl = "https://drive.google.com/drive/folders/1mgWMNrL24N92irRgrd4WUL3hhYrfSsIj?usp=sharing";
        const meetingPdfDriveMap = {pdf_map_str};

        function getMeetingPdfUrl(meetingNumber) {{
            if (!meetingNumber) return gdriveMainFolderUrl;
            let numStr = meetingNumber.toString().replace(/[০-৯]/g, d => banglaDigits[d]).replace(/\\D/g, '');
            if (numStr && meetingPdfDriveMap[numStr]) {{
                return meetingPdfDriveMap[numStr];
            }}
            return gdriveMainFolderUrl;
        }}
"""

start_idx = html.find('const masterDb = {')
end_idx = html.find('const banglaDigits =', start_idx)

new_text = f'const masterDb = {json_str};\n{pdf_inject_code}\n        '
new_html = html[:start_idx] + new_text + html[end_idx:]

with open(html_path, 'w', encoding='utf-8') as f:
    f.write(new_html)

# Regenerate Word Docs
print("📝 4/4 Updating MS Word Documents (cm word copy)...")
os.makedirs('cm word copy', exist_ok=True)
for f in os.listdir('cm word copy'):
    if f.endswith('.docx'):
        try:
            os.remove(os.path.join('cm word copy', f))
        except Exception:
            pass

for meeting in sorted_meetings:
    meeting_no = meeting.get('meeting_number', 'Unknown')
    safe_filename = re.sub(r'[\\/*?:"<>|]', '_', meeting_no)
    filename = f"cm word copy/{safe_filename}.docx"
    
    doc = Document()
    doc.add_heading(f"কমিশন সভা: {meeting_no}", 0)
    doc.add_paragraph(f"তারিখ: {meeting.get('meeting_date', '')}")
    
    for a in meeting.get('agendas', []):
        agenda_no = a.get('agenda_no', '')
        doc.add_heading(f"এজেন্ডা {agenda_no}", level=1)
        
        doc.add_heading("বিষয় (Subject):", level=2)
        doc.add_paragraph(a.get('subject', ''))
        
        doc.add_heading("সিদ্ধান্ত (Decision):", level=2)
        doc.add_paragraph(a.get('decision', ''))
        
        if a.get('fine_amount'):
            doc.add_paragraph(f"প্রশাসনিক জরিমানা: {a.get('fine_amount')} টাকা")
            
        doc.add_paragraph("-" * 40)
        
    doc.save(filename)

print("🎉 DONE! Excel is now the Master Source. Web Portal, JSON Database, and Word files are all synchronized!")
